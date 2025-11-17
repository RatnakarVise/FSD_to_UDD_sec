import os
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ------------------------------------------------------------
# Utilities
# ------------------------------------------------------------

def is_subsection_heading(text: str) -> bool:
    """
    Detect headings like:
    3.1 Something
    4.2. Technical Architecture
    8.6. Exception Handling
    """
    return bool(re.match(r"^\d+\.\d+(\.|)\s", text.strip()))


def add_heading(doc, text, level):
    """
    Heading LEVEL 1 only (TOC included).
    LEVEL 0 = TITLE.
    """
    para = doc.add_heading(text, level=level)
    return para if para else doc.add_paragraph(text, style="Heading1")


def add_normal_paragraph(doc, text):
    """
    ALWAYS create plain paragraph (no auto heading).
    Hard-disable Word automatic outline detection.
    """
    p = doc.add_paragraph()
    p.style = 'Normal'

    run = p.add_run(text)
    run.bold = False

    # Remove outline level so Word never treats it as heading
    p_props = p._p.get_or_add_pPr()
    outline = p_props.find(qn('w:outlineLvl'))
    if outline is not None:
        p_props.remove(outline)

    return p


# ------------------------------------------------------------
# Tables
# ------------------------------------------------------------

def add_table(doc, colnames, rows):
    table = doc.add_table(rows=1, cols=len(colnames))
    table.style = "Light List"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(colnames):
        hdr_cells[i].text = str(col)

    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val) if val is not None else ""

    return table


# ------------------------------------------------------------
# TOC + Bookmarks
# ------------------------------------------------------------

def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.text = 'TOC \\o "1" \\h \\z \\u'  # ONLY HEADING 1

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_bookmark(paragraph, name, bid):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_pageref_field(paragraph, bookmark_name):
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f"PAGEREF {bookmark_name} \\h"

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    result = OxmlElement("w:t")
    result.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(result)
    run._r.append(fld_end)


# ------------------------------------------------------------
# Parsing Logic
# ------------------------------------------------------------

def find_section_content(content_list, section_title):
    for sec in content_list:
        if sec.get('section_name', '').lower().strip() == section_title.lower().strip():
            return sec['content']
    return None


def find_all_table_like_chunks(text):
    if not text or not text.strip():
        return []

    lines = text.splitlines()
    chunks = []

    i = 0
    while i < len(lines):
        l = lines[i]

        # TABLE
        if l.count('|') >= 1 and (i + 1 < len(lines) and lines[i + 1].count('|') >= 1):
            buf = [l]
            i += 1
            while i < len(lines) and lines[i].count('|') >= 1:
                buf.append(lines[i])
                i += 1
            chunks.append(('table', "\n".join(buf).strip()))
            continue

        # TEXT
        if l.strip():
            chunks.append(('text', l.strip()))

        i += 1

    return chunks


def parse_markdown_table(table_md):
    lines = [l.strip() for l in table_md.splitlines() if l.strip()]
    if len(lines) < 2:
        return None, None

    if not lines[0].startswith('|'):
        return None, None

    rows = [[c.strip() for c in l.strip('|').split('|')] for l in lines]
    header = rows[0]
    data_rows = rows[1:]

    if all(len(r) == len(header) for r in data_rows):
        return header, data_rows

    return None, None


# ------------------------------------------------------------
# Build Full Document
# ------------------------------------------------------------

def build_document(content, sections, flow_diagram_agent=None, diagram_dir="diagrams"):
    doc = Document()

    # Title
    add_heading(doc, "Uniform Design Document", 0)

    # INDEX (TOC)
    idx = add_heading(doc, "Index", 1)
    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    # Manual Index Entries
    for i, section in enumerate(sections):
        p = doc.add_paragraph()
        title = section.get("title")
        bookmark_name = f"sec_{i+1}"

        run = p.add_run(f"{i+1}. {title}")
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )

        run.add_tab()
        add_pageref_field(p, bookmark_name)

    # SECTION CONTENT
    for i, section in enumerate(sections):
        title = section.get("title")
        bookmark_name = f"sec_{i+1}"

        # Heading 1 → included in TOC
        h = add_heading(doc, f"{i+1}. {title}", 1)
        add_bookmark(h, bookmark_name, i + 1)

        sec_content = find_section_content(content, title)

        # Parse into chunks
        chunks = find_all_table_like_chunks(sec_content or "")

        for typ, value in chunks:
            if typ == "text":

                # Subsection
                if is_subsection_heading(value):
                    p = doc.add_paragraph()
                    run = p.add_run(value)
                    run.bold = True

                    # FORCE Word to NEVER treat this as a heading
                    pPr = p._p.get_or_add_pPr()

                    # Set fake outline level = 9 (body text)
                    outline = OxmlElement('w:outlineLvl')
                    outline.set(qn('w:val'), '9')
                    pPr.append(outline)

                    # Remove auto-outline behavior
                    autoOutline = OxmlElement('w:autoOutline')
                    pPr.append(autoOutline)

                    # Disable numbering recognition
                    numPr = OxmlElement('w:numPr')
                    ilvl = OxmlElement('w:ilvl')
                    ilvl.set(qn('w:val'), '9')
                    numId = OxmlElement('w:numId')
                    numId.set(qn('w:val'), '0')
                    numPr.append(ilvl)
                    numPr.append(numId)
                    pPr.append(numPr)

                    continue


                # Normal text
                add_normal_paragraph(doc, value)

            elif typ == "table":
                colnames, rows = parse_markdown_table(value)
                if colnames and rows:
                    add_table(doc, colnames, rows)
                else:
                    add_normal_paragraph(doc, value)

    doc.add_paragraph("\nDocument generated by PWC AI-powered ABAP UDD Spec Assistant.")
    return doc
